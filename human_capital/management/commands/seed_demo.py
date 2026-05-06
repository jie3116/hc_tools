from __future__ import annotations

from pathlib import Path

from django.core.files import File
from django.core.management.base import BaseCommand
from docx import Document

from human_capital.models import Employee, LetterTemplate, PolicyDocument


class Command(BaseCommand):
    help = "Seed demo data untuk human capital app."

    def handle(self, *args, **options):
        Employee.objects.get_or_create(
            employee_code="EMP-001",
            defaults={
                "full_name": "Budi Santoso",
                "email": "budi@company.com",
                "department": "Engineering",
                "position": "Software Engineer",
                "employment_status": "Permanent",
                "join_date": "2023-02-01",
                "manager_name": "Rina Hartati",
            },
        )
        Employee.objects.get_or_create(
            employee_code="EMP-002",
            defaults={
                "full_name": "Siti Lestari",
                "email": "siti@company.com",
                "department": "Finance",
                "position": "Finance Analyst",
                "employment_status": "Permanent",
                "join_date": "2022-07-15",
                "manager_name": "Arif Prabowo",
            },
        )
        PolicyDocument.objects.get_or_create(
            title="Kebijakan Surat Keterangan Kerja",
            language="ID",
            defaults={
                "source_name": "kebijakan_surat_kerja.txt",
                "content": "Karyawan dapat mengajukan surat keterangan kerja melalui aplikasi HC. Admin HC memverifikasi kelengkapan data dan memilih template. Setelah lolos review admin, approver memberikan persetujuan akhir. Dokumen yang sudah disetujui dapat diunduh oleh karyawan.",
            },
        )
        PolicyDocument.objects.get_or_create(
            title="Annual Leave Policy",
            language="EN",
            defaults={
                "source_name": "annual_leave_policy.txt",
                "content": "Permanent employees are entitled to 12 days of annual leave after completing 12 months of service. Leave requests must be submitted through the HC portal and approved by the direct manager.",
            },
        )
        self._ensure_template("ID", "Template Surat Keterangan Kerja", "template_id.docx")
        self._ensure_template("EN", "Certificate of Employment Template", "template_en.docx")
        self.stdout.write(self.style.SUCCESS("Demo data seeded."))

    def _ensure_template(self, language: str, name: str, file_name: str) -> None:
        if LetterTemplate.objects.filter(name=name, language=language).exists():
            return
        temp_dir = Path("media") / "seed"
        temp_dir.mkdir(parents=True, exist_ok=True)
        path = temp_dir / file_name
        document = Document()
        if language == "ID":
            document.add_heading("Surat Keterangan Kerja", level=1)
            document.add_paragraph("Nomor: {{request_no}}")
            document.add_paragraph("Nama: {{full_name}}")
            document.add_paragraph("NIK: {{employee_code}}")
            document.add_paragraph("Departemen: {{department}}")
            document.add_paragraph("Jabatan: {{position}}")
            document.add_paragraph("Keperluan: {{purpose}}")
            document.add_paragraph("Ditujukan kepada: {{recipient_name}}")
            document.add_paragraph("Tanggal terbit: {{today}}")
        else:
            document.add_heading("Certificate of Employment", level=1)
            document.add_paragraph("Reference No: {{request_no}}")
            document.add_paragraph("Name: {{full_name}}")
            document.add_paragraph("Employee ID: {{employee_code}}")
            document.add_paragraph("Department: {{department}}")
            document.add_paragraph("Position: {{position}}")
            document.add_paragraph("Purpose: {{purpose}}")
            document.add_paragraph("Recipient: {{recipient_name}}")
            document.add_paragraph("Issue date: {{today}}")
        document.save(path)
        with path.open("rb") as stream:
            template = LetterTemplate(name=name, language=language)
            template.file.save(file_name, File(stream), save=True)
