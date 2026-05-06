from __future__ import annotations

import json
import shutil
import tempfile
from io import BytesIO

from django.contrib.auth import get_user_model
from django.contrib.auth.models import Group
from django.core.management import call_command
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import Client, TestCase, override_settings
from docx import Document

from human_capital.access import ROLE_ADMIN_HC, ROLE_APPROVER, ROLE_EMPLOYEE
from human_capital.models import Employee, LetterRequest, LetterTemplate
from human_capital.services import (
    WorkflowError,
    admin_review_letter_request,
    approver_review_letter_request,
    generate_request_no,
)

User = get_user_model()


class HumanCapitalBaseTestCase(TestCase):
    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        cls._temp_media = tempfile.mkdtemp()
        cls._override = override_settings(MEDIA_ROOT=cls._temp_media)
        cls._override.enable()

    @classmethod
    def tearDownClass(cls):
        cls._override.disable()
        shutil.rmtree(cls._temp_media, ignore_errors=True)
        super().tearDownClass()

    @classmethod
    def setUpTestData(cls):
        for role_name in (ROLE_EMPLOYEE, ROLE_ADMIN_HC, ROLE_APPROVER):
            Group.objects.get_or_create(name=role_name)

        cls.employee_user = User.objects.create_user(username="employee", password="password123")
        cls.admin_user = User.objects.create_user(username="admin", password="password123")
        cls.approver_user = User.objects.create_user(username="approver", password="password123")

        cls.employee_user.groups.add(Group.objects.get(name=ROLE_EMPLOYEE))
        cls.admin_user.groups.add(Group.objects.get(name=ROLE_ADMIN_HC))
        cls.approver_user.groups.add(Group.objects.get(name=ROLE_APPROVER))

        cls.employee_profile = Employee.objects.create(
            user=cls.employee_user,
            employee_code="EMP-001",
            full_name="Budi Santoso",
            email="budi@example.com",
            department="Engineering",
            position="Software Engineer",
            employment_status="Permanent",
            join_date="2024-01-01",
            manager_name="Rina",
        )
        cls.other_employee = Employee.objects.create(
            employee_code="EMP-002",
            full_name="Siti Lestari",
            email="siti@example.com",
            department="Finance",
            position="Analyst",
            employment_status="Permanent",
            join_date="2024-02-01",
            manager_name="Arif",
        )
        cls.template_id = cls._create_template("ID Template", "ID")

    @classmethod
    def _create_template(cls, name: str, language: str) -> LetterTemplate:
        document = Document()
        document.add_paragraph("Nomor: {{request_no}}")
        document.add_paragraph("Nama: {{full_name}}")
        document.add_paragraph("Tujuan: {{recipient_name}}")
        document.add_paragraph("Tanggal: {{today}}")
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        uploaded_file = SimpleUploadedFile(
            f"{name.lower().replace(' ', '_')}.docx",
            buffer.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        return LetterTemplate.objects.create(name=name, language=language, file=uploaded_file)

    def login(self, user):
        client = Client()
        assert client.login(username=user.username, password="password123")
        return client


class PermissionTests(HumanCapitalBaseTestCase):
    def test_dashboard_requires_login(self):
        response = self.client.get("/")
        self.assertEqual(response.status_code, 302)
        self.assertIn("/accounts/login/", response["Location"])

    def test_employee_cannot_access_employee_master_data(self):
        client = self.login(self.employee_user)
        response = client.get("/employees/")
        self.assertEqual(response.status_code, 403)

    def test_employee_only_sees_own_letter_requests(self):
        own_request = LetterRequest.objects.create(
            request_no="REQ-OWN",
            employee=self.employee_profile,
            template=self.template_id,
            purpose="Visa",
            recipient_name="Embassy",
            language="ID",
        )
        LetterRequest.objects.create(
            request_no="REQ-OTHER",
            employee=self.other_employee,
            template=self.template_id,
            purpose="Bank",
            recipient_name="Bank",
            language="ID",
        )

        client = self.login(self.employee_user)

        web_response = client.get("/letters/")
        self.assertContains(web_response, own_request.request_no)
        self.assertNotContains(web_response, "REQ-OTHER")

        api_response = client.get("/api/requests/")
        payload = json.loads(api_response.content)
        self.assertEqual(api_response.status_code, 200)
        self.assertEqual(len(payload), 1)
        self.assertEqual(payload[0]["request_no"], own_request.request_no)


class WorkflowTests(HumanCapitalBaseTestCase):
    def test_generate_request_no_is_unique(self):
        numbers = {generate_request_no() for _ in range(100)}
        self.assertEqual(len(numbers), 100)

    def test_employee_can_submit_request_for_own_profile(self):
        client = self.login(self.employee_user)
        response = client.post(
            "/letters/new/",
            {
                "template": self.template_id.pk,
                "purpose": "Pengajuan visa",
                "recipient_name": "Kedutaan",
                "language": "ID",
                "notes": "Perjalanan dinas",
            },
        )
        self.assertEqual(response.status_code, 302)
        created = LetterRequest.objects.get()
        self.assertEqual(created.employee, self.employee_profile)
        self.assertEqual(created.status, LetterRequest.STATUS_SUBMITTED)

    def test_api_assigns_employee_profile_for_employee_user(self):
        client = self.login(self.employee_user)
        response = client.post(
            "/api/requests/",
            data={
                "template": self.template_id.pk,
                "purpose": "Pengajuan visa",
                "recipient_name": "Kedutaan",
                "language": "ID",
                "notes": "Perjalanan dinas",
            },
        )
        self.assertEqual(response.status_code, 201)
        created = LetterRequest.objects.get()
        self.assertEqual(created.employee, self.employee_profile)

    def test_workflow_guards_and_document_generation(self):
        template = self._create_template("Workflow Template", "ID")
        request = LetterRequest.objects.create(
            request_no="REQ-001",
            employee=self.employee_profile,
            template=template,
            purpose="Visa",
            recipient_name="Embassy",
            language="ID",
            status=LetterRequest.STATUS_SUBMITTED,
        )

        with self.assertRaises(WorkflowError):
            approver_review_letter_request(request, True, "too early")

        admin_review_letter_request(request, True, "ok")
        request.refresh_from_db()
        self.assertEqual(request.status, LetterRequest.STATUS_PENDING_APPROVAL)

        with self.assertRaises(WorkflowError):
            admin_review_letter_request(request, True, "second review")

        approver_review_letter_request(request, True, "approved")
        request.refresh_from_db()
        self.assertEqual(request.status, LetterRequest.STATUS_APPROVED)
        self.assertIsNotNone(request.generated_file)
        self.assertTrue(request.generated_file.name.endswith(".docx"))


class OnboardingCommandTests(HumanCapitalBaseTestCase):
    def test_seed_demo_users_creates_roles_and_links_employee(self):
        Employee.objects.filter(pk=self.employee_profile.pk).update(user=None)

        call_command("seed_demo_users")

        admin_user = User.objects.get(username="hc-admin")
        approver_user = User.objects.get(username="hc-approver")
        employee_user = User.objects.get(username="hc-employee")
        employee = Employee.objects.get(employee_code="EMP-001")

        self.assertTrue(admin_user.groups.filter(name=ROLE_ADMIN_HC).exists())
        self.assertTrue(approver_user.groups.filter(name=ROLE_APPROVER).exists())
        self.assertTrue(employee_user.groups.filter(name=ROLE_EMPLOYEE).exists())
        self.assertEqual(employee.user, employee_user)
